"""
Requirements:
python-docx
git+https://github.com/ssuwani/pytube

Takes in a video url
extracts the caption track (language)
extracts the description
extracts the timestamps from the description
Intersplices the
"""
from datetime import datetime
import re
import os
from collections import defaultdict
from pytube import YouTube
from docx import Document



class Video:
    """wrapper for containing video information"""
    def __init__(self, url):
        self._video = YouTube(url)
        self._captionstr = self._video.caption_tracks[0].generate_srt_captions()
        self._descriptionstr = self._video.description

    def get_caption_track(self):
        return CaptionTrack(self._captionstr, self._descriptionstr)

    def get_caption_document(self):
        caption_track = self.get_caption_track()
        caption_dict = caption_track.as_dict()
        document = CaptionDocument(caption_dict,self._video.title)
        return document

    def get_caption_dict(self):
        caption_track = self.get_caption_track()
        caption_dict = caption_track.as_dict()
        return caption_dict


class Caption:
    """represents an individual caption"""
    def __init__(self, captionstr):
        caption_lines = captionstr.splitlines()
        self._idstr= caption_lines[0]
        self._time_str = caption_lines[1]
        self._text = caption_lines[2]

    def __str__(self):
        """Returns the caption text"""
        return self._text

    @property
    def _start_end_timestr(self):
        times = self._time_str.split(' --> ')
        return times[0], times[1]

    @property
    def start_time(self):
        start_timestr = self._start_end_timestr[0]
        time = datetime.strptime(start_timestr, "%H:%M:%S,%f")
        return time

    @property
    def end_time(self):
        start_timestr = self._start_end_timestr[0]
        time = datetime.strptime(start_timestr, "%H:%M:%S,%f")
        return time

    @property
    def id(self):
        return int(self._idstr)

    @property
    def text(self):
        return self._text

    @text.setter
    def text(self, value):
        self._text = value


class Description:
    """Represents a description. outputs summary and timestamps"""
    def __init__(self, descriptionstr):
        self._description_str = descriptionstr
        self._description_lines = descriptionstr.split("\n")

    @property
    def timestamps(self):
        timestamps = list()
        for line in self._description_lines:
            try:
                result = re.search("[0-9]+[:][0-9]+[:]+[0-9]+", line)
                time_count = datetime.strptime(result.group(), '%H:%M:%S')
                timestamp = self._timestamp_from_re(line, result, time_count)
                timestamps.append(timestamp)
            except:
                try:
                    result = re.search("[0-9]+[:][0-9]+", line)
                    time_count = datetime.strptime(result.group(), '%M:%S')
                    timestamp = self._timestamp_from_re(line, result, time_count)
                    timestamps.append(timestamp)
                except:
                    continue

        return timestamps

    @staticmethod
    def _timestamp_from_re(timestampstr, result, time_count):
        chap_name = timestampstr.replace(result.group(0), "").rstrip(' :\n')
        return Timestamp(time_count, chap_name)


class Timestamp:
    """Represents a timestamp. can check to see if a caption is before or after it"""
    def __init__(self, position, name):
        self.position = position
        self.name = name

    def before_caption(self, caption):
        """if before caption, true"""
        return self.position < caption.start_time

    def after_caption(self, caption):
        return self.position > caption.end_time


class CaptionTrack:
    """Contains all the captions, has methods for extracting the captions into something coherent"""
    def __init__(self, captionstr, descriptionstr):

        self._description = Description(descriptionstr)

        captionstrs = captionstr.split('\n\n')  # splits into each group

        self.captions = [Caption(captionstr) for captionstr in captionstrs]
        self.captions = sorted(self.captions, key=lambda x: x.id)

    @property
    def timestamps(self):
        return self._description.timestamps

    def as_list(self):
        caption_strs = list()
        for caption in self.captions:
            if not caption.text.endswith(' '):
                caption.text = caption.text + ' '
            caption_strs.append(caption.text)
        return caption_strs

    def as_text(self):
        return ''.join(self.as_list())

    def as_lines(self):
        return '\n'.join(self.as_list())

    def as_dict(self):
        result = defaultdict(list)




        ts_id = 0
        ts_max = len(self.timestamps)
        cap_id = 0
        cap_max = len(self.captions)
        while True:
            if ts_id == ts_max or cap_id == cap_max:
                break

            prev_ts_id = ts_id - 1 if ts_id > 0 else 0
            cap = self.captions[cap_id]
            ts = self.timestamps[ts_id]
            if ts.position > cap.start_time:
                prev_ts = self.timestamps[prev_ts_id]
                key = (prev_ts.position, prev_ts.name)
                result[key].append(cap)
                cap_id += 1
            else:
                ts_id += 1
        return result


class CaptionParagraph:
    def __init__(self, title, captions):
        self._title = title
        self._captions = captions

    @property
    def heading(self):
        return self._title.title()

    @property
    def body(self):
        text = [caption.text for caption in self._captions]
        text = [line + ' ' if not line.endswith(' ') else line for line in text]
        return ''.join(text)


class CaptionDocument:
    _invalid_save_chars = ['|']

    def __init__(self, caption_dict, title, section=None, path=None):
        self.paragraphs = [CaptionParagraph(k[1], v) for k, v in caption_dict.items()]
        self.section = section
        self.title = title
        self._doc = Document(path)
        self._path = None

        for paragraph in self.paragraphs:
            self.add_paragraph(paragraph)

    def add_paragraph(self, paragraph):
        self.add_heading(paragraph.heading)
        self.add_body_text(paragraph.body)

    def add_heading(self, text=''):
        p = self._doc.add_paragraph()
        p.text = text
        p.style = 'Heading 1'

    def add_body_text(self, text=''):
        p = self._doc.add_paragraph()
        p.text = text

    def save_to(self, dirpath):
        path = os.path.join(dirpath, self.title)
        self.save_as(path)

    def _as_valid_save_path(self, path):
        path = path + '.docx' if not path.endswith('.docx') else path
        for char in self._invalid_save_chars:
            path = path.replace(char, '-')
        path = os.path.abspath(path)

        return path

    def save_as(self, path):
        path = self._as_valid_save_path(path)
        self._doc.save(path)


if __name__ == '__main__':

    path = r"https://www.youtube.com/watch?v=GqPGXG5TlZw"
    vid = Video(path)
    doc = vid.get_caption_document()
    doc.save_to('docs')